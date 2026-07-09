from __future__ import annotations

import json
import sys
from pathlib import Path
from tempfile import TemporaryDirectory

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from openpyxl import Workbook  # noqa: E402

from app.services.chunker import chunk_document, summarize_chunk_quality  # noqa: E402
from app.services.document_parser import parse_document  # noqa: E402


def write_report_md(path: Path) -> None:
    path.write_text(
        """# 数据中心行业研究报告（2024）

## 一、行业概览

数据中心是数字经济的物理底座。香港市场受土地供应、电力约束和跨境数据合规三重因素影响，新建项目资本开支普遍高于内地同类项目 15%-25%。

## 二、核心风险

### 2.1 电力与能耗

PUE 约束趋严，备用电源配置成本上升。若上架率低于 70%，单位收入能耗成本将显著侵蚀 EBITDA 利润率。

### 2.2 客户需求

超大规模云厂商租约通常附带价格下调条款。合同期内若客户迁移或缩容，空置率风险将直接传导至现金流预测。

## 三、结论

综合判断，数据中心项目应重点关注电力供给稳定性、长期租约质量与客户集中度。香港项目还需额外评估土地续期与合规成本。
""",
        encoding="utf-8",
    )


def write_report_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("商业地产行业概览", level=1)
    doc.add_paragraph(
        "2024年香港甲级写字楼空置率维持在 12.8% 左右，租金较峰值回落约 18%。"
        "核心区与次级市场分化加剧，租户议价能力增强。"
    )
    doc.add_heading("供需格局", level=2)
    doc.add_paragraph(
        "新增供应集中在启德与东九龙，未来三年预计新增面积约 280 万平方呎。"
        "跨国金融机构持续优化办公面积，需求端以成本压缩和灵活办公为主。"
    )
    doc.add_heading("投资要点", level=2)
    doc.add_paragraph(
        "资本化率上行压制资产估值，交易活跃度下降。"
        "具备稳定租约、低杠杆和优质区位的大宗资产仍具配置价值。"
    )
    table = doc.add_table(rows=4, cols=3)
    headers = ["区域", "空置率", "租金同比"]
    for col, value in enumerate(headers):
        table.rows[0].cells[col].text = value
    rows = [
        ("中环", "9.5%", "-6%"),
        ("金钟", "10.2%", "-5%"),
        ("九龙东", "15.1%", "-9%"),
    ]
    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            table.rows[row_index].cells[col_index].text = value
    doc.save(path)


def write_financial_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "财务摘要"
    ws.append(["指标", "2022", "2023", "2024E"])
    metrics = [
        ("营业收入（亿港元）", "128.4", "141.2", "156.8"),
        ("EBITDA 利润率", "32.1%", "30.8%", "29.5%"),
        ("净负债/EBITDA", "3.8x", "3.5x", "3.2x"),
        ("资本开支", "22.6", "25.1", "28.4"),
        ("自由现金流", "18.3", "16.9", "15.2"),
    ]
    for row in metrics:
        ws.append(list(row))
    ws2 = wb.create_sheet("区域收入")
    ws2.append(["区域", "收入占比", "同比增速"])
    for region, share, growth in [
        ("香港", "46%", "4.2%"),
        ("大湾区", "31%", "11.5%"),
        ("东南亚", "23%", "8.7%"),
    ]:
        ws2.append([region, share, growth])
    wb.save(path)


def chunk_length_histogram(lengths: list[int]) -> dict[str, int]:
    buckets = {"<180": 0, "180-500": 0, "500-820": 0, "820-1200": 0, "1200-1700": 0, ">1700": 0}
    for length in lengths:
        if length < 180:
            buckets["<180"] += 1
        elif length < 500:
            buckets["180-500"] += 1
        elif length < 820:
            buckets["500-820"] += 1
        elif length < 1200:
            buckets["820-1200"] += 1
        elif length < 1700:
            buckets["1200-1700"] += 1
        else:
            buckets[">1700"] += 1
    return buckets


def analyze_file(label: str, path: Path) -> dict:
    parsed = parse_document(path)
    chunks = chunk_document(parsed, document_id=f"sample_{label}")
    quality = summarize_chunk_quality(chunks)
    lengths = [len(chunk.text) for chunk in chunks]
    parent_lengths = [len(str(chunk.metadata.get("parent_text") or "")) for chunk in chunks]
    block_types = {}
    for chunk in chunks:
        block_type = str(chunk.metadata.get("block_type") or "unknown")
        block_types[block_type] = block_types.get(block_type, 0) + 1

    samples = []
    for chunk in chunks[:3]:
        samples.append(
            {
                "id": chunk.id,
                "chars": len(chunk.text),
                "parent_chars": len(str(chunk.metadata.get("parent_text") or "")),
                "heading_path": chunk.metadata.get("heading_path"),
                "source_locator": chunk.source_locator,
                "block_type": chunk.metadata.get("block_type"),
                "parent_differs": chunk.metadata.get("parent_differs_from_child"),
                "text_preview": chunk.text[:120].replace("\n", " "),
            }
        )

    return {
        "label": label,
        "file": path.name,
        "parsed_blocks": len(parsed.blocks),
        "block_types_in_parser": _count_parser_blocks(parsed),
        "chunk_count": len(chunks),
        "quality": quality,
        "length_histogram": chunk_length_histogram(lengths),
        "parent_length_histogram": chunk_length_histogram(parent_lengths),
        "chunk_block_types": block_types,
        "samples": samples,
    }


def _count_parser_blocks(parsed) -> dict[str, int]:
    counts: dict[str, int] = {}
    for block in parsed.blocks:
        counts[block.block_type] = counts.get(block.block_type, 0) + 1
    return counts


def main() -> int:
    with TemporaryDirectory(prefix="chunk_samples_") as tmp:
        tmp_path = Path(tmp)
        files = {
            "md_report": tmp_path / "data_center_report.md",
            "docx_report": tmp_path / "office_sector_report.docx",
            "xlsx_financial": tmp_path / "financial_summary.xlsx",
        }
        write_report_md(files["md_report"])
        write_report_docx(files["docx_report"])
        write_financial_xlsx(files["xlsx_financial"])

        results = [analyze_file(label, path) for label, path in files.items()]
        print(json.dumps(results, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
