from __future__ import annotations

import logging
import re
from bisect import bisect_right
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app.config import get_settings

logger = logging.getLogger(__name__)


_STEP_SPLIT_RE = re.compile(r"(?=第[一二三四五六七八九十\d]+步[：:])")
@dataclass(frozen=True)
class ParsedBlock:
    text: str
    source_locator: str
    block_type: str = "text"
    page_start: int | None = None
    page_end: int | None = None
    heading_path: tuple[str, ...] = field(default_factory=tuple)


@dataclass(frozen=True)
class ParsedDocument:
    blocks: list[ParsedBlock]

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks if block.text.strip())


class DocumentParserError(RuntimeError):
    pass


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        document = _parse_text(path)
    elif suffix == ".pdf":
        document = _parse_pdf(path)
    elif suffix == ".docx":
        document = _parse_docx(path)
    elif suffix == ".xlsx":
        document = _parse_xlsx(path)
    else:
        raise DocumentParserError(f"Unsupported file type: {suffix or 'unknown'}")
    return ParsedDocument(blocks=normalize_blocks(document.blocks))


def normalize_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    settings = get_settings()
    normalized = [_normalize_block_artifacts(block) for block in blocks]
    expanded: list[ParsedBlock] = []
    for block in normalized:
        expanded.extend(_expand_block(block))
    split_mixed = _split_mixed_content_blocks(expanded)
    merged = _merge_short_text_blocks(split_mixed, settings.chunk_child_min_chars, settings.chunk_child_target_chars)
    return _promote_table_like_blocks(merged)


def _normalize_block_artifacts(block: ParsedBlock) -> ParsedBlock:
    if block.page_start is None:
        return block
    return ParsedBlock(
        text=_normalize_pdf_artifacts(block.text),
        source_locator=block.source_locator,
        block_type=block.block_type,
        page_start=block.page_start,
        page_end=block.page_end,
        heading_path=block.heading_path,
    )


def _normalize_pdf_artifacts(text: str) -> str:
    cleaned = text
    cleaned = re.sub(r"(\d)\.\s+(\d)", r"\1.\2", cleaned)
    cleaned = re.sub(r"(\d)\s+%", r"\1%", cleaned)
    cleaned = re.sub(r"([。！？；;])\s+", r"\1", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip()


def _split_mixed_content_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    split_blocks: list[ParsedBlock] = []
    for block in blocks:
        if block.block_type != "text":
            split_blocks.append(block)
            continue
        split_blocks.extend(_split_text_and_table_runs(block))
    return split_blocks


def _split_text_and_table_runs(block: ParsedBlock) -> list[ParsedBlock]:
    lines = [line.strip() for line in block.text.splitlines() if line.strip()]
    if len(lines) < 4:
        return [block]

    runs: list[tuple[str, list[str]]] = []
    current_type: str | None = None
    current_lines: list[str] = []
    for line in lines:
        line_type = "table" if _looks_like_table_row(line) else "text"
        if current_type != line_type:
            if current_lines and current_type is not None:
                runs.append((current_type, current_lines))
            current_type = line_type
            current_lines = [line]
        else:
            current_lines.append(line)
    if current_lines and current_type is not None:
        runs.append((current_type, current_lines))

    if len(runs) <= 1:
        return [block]

    split_blocks: list[ParsedBlock] = []
    for run_type, run_lines in runs:
        if not run_lines:
            continue
        block_type = "table" if run_type == "table" and len(run_lines) >= 2 else "text"
        split_blocks.append(
            ParsedBlock(
                text="\n".join(run_lines),
                source_locator=block.source_locator,
                block_type=block_type,
                page_start=block.page_start,
                page_end=block.page_end,
                heading_path=block.heading_path,
            )
        )
    return split_blocks or [block]


def _expand_block(block: ParsedBlock) -> list[ParsedBlock]:
    if block.block_type in {"heading", "table"}:
        return [block]
    if block.block_type != "text":
        return [block]

    sections = _split_glued_step_sections(block.text)
    if len(sections) <= 1:
        return [block]

    expanded: list[ParsedBlock] = []
    current_heading = list(block.heading_path)
    for section in sections:
        heading = _step_heading(section)
        if heading:
            current_heading = _update_heading_path(current_heading, heading[0], heading[1])
            expanded.append(
                ParsedBlock(
                    text=heading[1],
                    source_locator=block.source_locator,
                    block_type="heading",
                    page_start=block.page_start,
                    page_end=block.page_end,
                    heading_path=tuple(current_heading),
                )
            )
            body = heading[2].strip()
            if body:
                expanded.append(
                    ParsedBlock(
                        text=body,
                        source_locator=block.source_locator,
                        block_type="text",
                        page_start=block.page_start,
                        page_end=block.page_end,
                        heading_path=tuple(current_heading),
                    )
                )
            continue
        expanded.append(
            ParsedBlock(
                text=section,
                source_locator=block.source_locator,
                block_type="text",
                page_start=block.page_start,
                page_end=block.page_end,
                heading_path=tuple(current_heading),
            )
        )
    return expanded or [block]


def _split_glued_step_sections(text: str) -> list[str]:
    parts = [part.strip() for part in _STEP_SPLIT_RE.split(text) if part.strip()]
    return parts if len(parts) > 1 else [text]


def _merge_short_text_blocks(blocks: list[ParsedBlock], min_chars: int, target_chars: int) -> list[ParsedBlock]:
    if not blocks:
        return blocks

    merged: list[ParsedBlock] = []
    buffer: ParsedBlock | None = None

    def flush() -> None:
        nonlocal buffer
        if buffer is not None:
            merged.append(buffer)
            buffer = None

    for block in blocks:
        if block.block_type != "text":
            flush()
            merged.append(block)
            continue

        if buffer is None:
            buffer = block
            continue

        combined_len = len(buffer.text) + len(block.text)
        same_heading = buffer.heading_path == block.heading_path
        same_locator = buffer.source_locator == block.source_locator
        should_merge = same_heading and same_locator and (
            len(buffer.text) < min_chars or combined_len <= target_chars
        )
        if should_merge:
            buffer = ParsedBlock(
                text=f"{buffer.text}\n\n{block.text}".strip(),
                source_locator=buffer.source_locator,
                block_type="text",
                page_start=_min_page_value(buffer.page_start, block.page_start),
                page_end=_max_page_value(buffer.page_end, block.page_end),
                heading_path=buffer.heading_path,
            )
            continue

        flush()
        buffer = block

    flush()
    return merged


def _promote_table_like_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    promoted: list[ParsedBlock] = []
    for block in blocks:
        if block.block_type != "text":
            promoted.append(block)
            continue
        lines = [line.strip() for line in block.text.splitlines() if line.strip()]
        if len(lines) < 3:
            promoted.append(block)
            continue
        table_like = sum(1 for line in lines if _looks_like_table_row(line))
        if table_like / len(lines) >= 0.6:
            promoted.append(
                ParsedBlock(
                    text=block.text,
                    source_locator=block.source_locator,
                    block_type="table",
                    page_start=block.page_start,
                    page_end=block.page_end,
                    heading_path=block.heading_path,
                )
            )
        else:
            promoted.append(block)
    return promoted


def _min_page_value(left: int | None, right: int | None) -> int | None:
    if left is None:
        return right
    if right is None:
        return left
    return min(left, right)


def _max_page_value(left: int | None, right: int | None) -> int | None:
    if left is None:
        return right
    if right is None:
        return left
    return max(left, right)


def _parse_text(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks: list[ParsedBlock] = []
    current_heading: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            blocks.append(
                ParsedBlock(
                    text="\n".join(buffer).strip(),
                    source_locator="全文",
                    block_type="text",
                    heading_path=tuple(current_heading),
                )
            )
            buffer.clear()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        heading = _markdown_heading(line) or _plain_heading(line)
        if heading:
            flush()
            current_heading = _update_heading_path(current_heading, heading[0], heading[1])
            blocks.append(
                ParsedBlock(
                    text=heading[1],
                    source_locator="全文",
                    block_type="heading",
                    heading_path=tuple(current_heading),
                )
            )
        else:
            buffer.append(line)
    flush()

    if not blocks:
        raise DocumentParserError("Text file contains no readable text.")
    return ParsedDocument(blocks=blocks)


def _parse_pdf(path: Path) -> ParsedDocument:
    reader = PdfReader(str(path))
    page_lines: list[list[str]] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        page_lines.append(_clean_lines(text))

    repeated = _repeated_margin_lines(page_lines)
    text_blocks: list[ParsedBlock] = []
    current_heading: list[str] = []
    for index, lines in enumerate(page_lines, start=1):
        clean_lines = [line for line in lines if line not in repeated]
        if not clean_lines:
            continue
        page_blocks, current_heading = _blocks_from_lines(
            clean_lines,
            source_locator=f"第 {index} 页",
            page=index,
            heading_path=current_heading,
        )
        text_blocks.extend(page_blocks)

    table_blocks = _extract_pdf_tables_with_pdfplumber(path)
    blocks = _merge_pdf_text_and_table_blocks(text_blocks, table_blocks)

    if not blocks:
        raise DocumentParserError("PDF contains no extractable text. OCR is not supported in this demo stage.")
    return ParsedDocument(blocks=blocks)


def _extract_pdf_tables_with_pdfplumber(path: Path) -> list[ParsedBlock]:
    table_blocks: list[ParsedBlock] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            last_caption = ""
            for page_index, page in enumerate(pdf.pages, start=1):
                try:
                    found_tables = page.find_tables() or []
                except Exception:
                    logger.warning("pdfplumber failed to extract tables on page %s of %s", page_index, path.name)
                    continue
                page_table_index = 0
                for found_table in found_tables:
                    matrix = _pdf_table_matrix(page, found_table)
                    markdown = _pdf_table_to_markdown(matrix)
                    if not markdown:
                        continue
                    page_table_index += 1
                    caption = _find_pdf_table_caption(page, found_table.bbox)
                    if caption:
                        last_caption = caption
                    elif found_table.bbox[1] < _PDF_TABLE_CONTINUATION_TOP and last_caption:
                        # A table starting at the very top of a page usually continues
                        # the previous page's table, so it inherits that caption.
                        caption = f"{last_caption}（续）"
                    locator_parts = [f"第 {page_index} 页", f"表格 {page_table_index}"]
                    if caption:
                        locator_parts.append(caption)
                    table_blocks.append(
                        ParsedBlock(
                            text=markdown,
                            source_locator=" / ".join(locator_parts),
                            block_type="table",
                            page_start=page_index,
                            page_end=page_index,
                            heading_path=(caption,) if caption else (),
                        )
                    )
    except Exception:
        logger.warning("pdfplumber failed to open %s; continuing with text-only PDF parse", path.name, exc_info=True)
    return table_blocks


_PDF_TABLE_CAPTION_RE = re.compile(r"^(图表|图|表|Table|Exhibit)\s*\d+")
_PDF_TABLE_CAPTION_SEARCH_HEIGHT = 46.0
_PDF_TABLE_CONTINUATION_TOP = 100.0


def _find_pdf_table_caption(page, bbox: tuple[float, float, float, float]) -> str:
    top = bbox[1]
    if top <= 4:
        return ""
    try:
        crop = page.crop((0, max(0.0, top - _PDF_TABLE_CAPTION_SEARCH_HEIGHT), page.width, top))
        text = crop.extract_text() or ""
    except Exception:
        return ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return ""
    candidate = lines[-1]
    if _PDF_TABLE_CAPTION_RE.match(candidate):
        return re.sub(r"\s+", " ", candidate)
    return ""


_PDF_TABLE_ROW_TOLERANCE = 5.0


def _pdf_table_matrix(page, found_table) -> list[list[str]] | None:
    """Rebuild the table from word coordinates using the detected column boundaries.

    pdfplumber's default lattice extraction can split logical rows when horizontal
    rules cut through them; word-level reconstruction keyed on the (reliable)
    vertical boundaries avoids losing rows.
    """
    try:
        boundaries = sorted({edge for cell in found_table.cells if cell for edge in (cell[0], cell[2])})
    except Exception:
        boundaries = []
    if len(boundaries) < 3:
        try:
            return found_table.extract()
        except Exception:
            return None

    try:
        crop = page.crop(found_table.bbox)
        words = crop.extract_words()
    except Exception:
        words = []
    if not words:
        try:
            return found_table.extract()
        except Exception:
            return None

    row_groups: list[dict] = []
    for word in sorted(words, key=lambda item: item["top"]):
        if row_groups and word["top"] - row_groups[-1]["top"] <= _PDF_TABLE_ROW_TOLERANCE:
            row_groups[-1]["words"].append(word)
        else:
            row_groups.append({"top": word["top"], "words": [word]})

    column_count = len(boundaries) - 1
    matrix: list[list[str]] = []
    for group in row_groups:
        cells = ["" for _ in range(column_count)]
        for word in sorted(group["words"], key=lambda item: item["x0"]):
            center = (word["x0"] + word["x1"]) / 2
            column = bisect_right(boundaries, center) - 1
            column = min(max(column, 0), column_count - 1)
            cells[column] = f"{cells[column]} {word['text']}".strip()
        matrix.append(cells)
    return matrix


def _pdf_table_to_markdown(raw_table: list[list] | None) -> str | None:
    if not raw_table:
        return None
    rows: list[str] = []
    for row in raw_table:
        if row is None:
            continue
        cells = [_normalize_pdf_table_cell(cell) for cell in row]
        cells = [cell for cell in cells if cell]
        if len(cells) >= 2:
            rows.append(" | ".join(cells))
    if len(rows) < 2:
        return None
    return "\n".join(rows)


def _normalize_pdf_table_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\n", " ").strip()
    return re.sub(r"\s+", " ", text)


def _table_token_set(text: str) -> set[str]:
    tokens = re.findall(r"[\u4e00-\u9fffA-Za-z0-9.%]+", text.lower())
    return {token for token in tokens if len(token) >= 2}


def _text_mostly_duplicates_table(text: str, table_tokens: set[str]) -> bool:
    if not table_tokens:
        return False
    text_tokens = _table_token_set(text)
    if len(text_tokens) < 6:
        return False
    overlap = text_tokens & table_tokens
    coverage = len(overlap) / max(1, len(text_tokens))
    table_coverage = len(overlap) / max(1, len(table_tokens))
    return coverage >= 0.72 and table_coverage >= 0.45


def _merge_pdf_text_and_table_blocks(
    text_blocks: list[ParsedBlock],
    table_blocks: list[ParsedBlock],
) -> list[ParsedBlock]:
    if not table_blocks:
        return text_blocks

    page_table_tokens: dict[int | None, set[str]] = {}
    for block in table_blocks:
        page_table_tokens.setdefault(block.page_start, set()).update(_table_token_set(block.text))

    filtered_text: list[ParsedBlock] = []
    for block in text_blocks:
        if (
            block.block_type == "text"
            and block.page_start in page_table_tokens
            and _text_mostly_duplicates_table(block.text, page_table_tokens[block.page_start])
        ):
            continue
        filtered_text.append(block)

    return filtered_text + table_blocks


def _parse_docx(path: Path) -> ParsedDocument:
    document = Document(str(path))
    blocks: list[ParsedBlock] = []
    current_heading: list[str] = []
    paragraph_index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph_index += 1
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        heading = _docx_heading(style_name, text) or _plain_heading(text)
        if heading:
            current_heading = _update_heading_path(current_heading, heading[0], heading[1])
            blocks.append(
                ParsedBlock(
                    text=heading[1],
                    source_locator=f"段落 {paragraph_index}",
                    block_type="heading",
                    heading_path=tuple(current_heading),
                )
            )
        else:
            blocks.append(
                ParsedBlock(
                    text=text,
                    source_locator=f"段落 {paragraph_index}",
                    block_type="text",
                    heading_path=tuple(current_heading),
                )
            )

    table_index = 0
    for table in document.tables:
        table_index += 1
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(
                ParsedBlock(
                    text="\n".join(rows),
                    source_locator=f"表格 {table_index}",
                    block_type="table",
                    heading_path=tuple(current_heading),
                )
            )

    if not blocks:
        raise DocumentParserError("DOCX contains no readable text.")
    return ParsedDocument(blocks=blocks)


def _parse_xlsx(path: Path) -> ParsedDocument:
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    blocks: list[ParsedBlock] = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))
        if rows:
            blocks.append(
                ParsedBlock(
                    text="\n".join(rows),
                    source_locator=f"工作表 {sheet.title}",
                    block_type="table",
                    heading_path=(sheet.title,),
                )
            )
    if not blocks:
        raise DocumentParserError("XLSX contains no readable cells.")
    return ParsedDocument(blocks=blocks)


def _blocks_from_lines(
    lines: list[str],
    source_locator: str,
    page: int,
    heading_path: list[str],
) -> tuple[list[ParsedBlock], list[str]]:
    blocks: list[ParsedBlock] = []
    buffer: list[str] = []
    current_heading = list(heading_path)

    def flush() -> None:
        if buffer:
            blocks.append(
                ParsedBlock(
                    text=_repair_wrapped_lines(buffer),
                    source_locator=source_locator,
                    block_type="text",
                    page_start=page,
                    page_end=page,
                    heading_path=tuple(current_heading),
                )
            )
            buffer.clear()

    for line in lines:
        heading = _plain_heading(line)
        if heading:
            flush()
            current_heading = _update_heading_path(current_heading, heading[0], heading[1])
            blocks.append(
                ParsedBlock(
                    text=heading[1],
                    source_locator=source_locator,
                    block_type="heading",
                    page_start=page,
                    page_end=page,
                    heading_path=tuple(current_heading),
                )
            )
        else:
            buffer.append(line)
    flush()
    return blocks, current_heading


def _clean_lines(text: str) -> list[str]:
    lines = []
    for line in text.splitlines():
        line = re.sub(r"\s+", " ", line).strip()
        if line:
            lines.append(line)
    return lines


def _clean_text(text: str) -> str:
    return "\n".join(_clean_lines(text))


def _repeated_margin_lines(page_lines: list[list[str]]) -> set[str]:
    candidates: list[str] = []
    for lines in page_lines:
        candidates.extend(lines[:3])
        candidates.extend(lines[-3:])
    counts = Counter(line for line in candidates if 3 <= len(line) <= 80)
    threshold = max(3, int(len(page_lines) * 0.35))
    repeated = {line for line, count in counts.items() if count >= threshold}
    return {line for line in repeated if not _plain_heading(line)}


def _repair_wrapped_lines(lines: list[str]) -> str:
    paragraphs: list[str] = []
    current = ""
    for line in lines:
        if _looks_like_table_row(line):
            if current:
                paragraphs.append(current)
                current = ""
            paragraphs.append(line)
            continue
        if not current:
            current = line
            continue
        if _should_join(current, line):
            current = f"{current}{line}"
        else:
            paragraphs.append(current)
            current = line
    if current:
        paragraphs.append(current)
    return "\n\n".join(paragraphs)


def _should_join(previous: str, current: str) -> bool:
    if previous.endswith(("。", "；", "：", ":", ".", "?", "？", "!", "！", "）", ")")):
        return False
    if re.match(r"^[-•·●▪]|^\d+[.)、]", current):
        return False
    if re.match(r"^[\d%+\-（(]", current) and not previous.endswith(("。", "！", "？", "；", "：", ":", ".", "!", ")", "）")):
        return True
    if re.search(r"[\u4e00-\u9fff]$", previous) and re.match(r"^[\u4e00-\u9fff（(]", current):
        if len(previous) >= 12:
            return True
    if len(previous) < 18:
        return False
    return True


def _looks_like_table_row(line: str) -> bool:
    if "|" in line:
        return True
    if re.search(r"\s{2,}", line):
        return True
    if re.match(r"^(指标|项目|区域|优点|缺点|情景|步骤|数据来源)", line):
        return True
    cells = re.split(r"\s{2,}", line.strip())
    return len(cells) >= 3 and all(len(cell) <= 40 for cell in cells)


def _markdown_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if not match:
        return None
    return len(match.group(1)), match.group(2).strip()


def _docx_heading(style_name: str, text: str) -> tuple[int, str] | None:
    match = re.match(r"heading\s+(\d+)", style_name)
    if match:
        return min(6, int(match.group(1))), text.strip()
    return None


def _plain_heading(line: str) -> tuple[int, str] | None:
    text = line.strip()
    if len(text) > 80:
        return None
    if _looks_like_table_row(text):
        return None
    patterns = [
        (r"^[一二三四五六七八九十]+[、.．]\s*(.+)$", 1),
        (r"^第[一二三四五六七八九十\d]+[章节部分步][：:]\s*(.+)$", 1),
        (r"^第[一二三四五六七八九十\d]+[章节部分步]\s+(.+)$", 1),
        (r"^\d+[、.．]\s+(.+)$", 2),
        (r"^[（(]\d+[）)]\s*(.+)$", 4),
    ]
    for pattern, level in patterns:
        match = re.match(pattern, text)
        if match:
            return level, text
    return None


def _step_heading(text: str) -> tuple[int, str, str] | None:
    match = re.match(r"^(第[一二三四五六七八九十\d]+步)[：:]\s*(.*)$", text.strip())
    if not match:
        return None
    title = match.group(1)
    body = match.group(2).strip()
    return 2, title, body


def _update_heading_path(current: list[str], level: int, title: str) -> list[str]:
    level = max(1, min(level, 6))
    next_path = current[: level - 1]
    next_path.append(title.strip())
    return next_path
